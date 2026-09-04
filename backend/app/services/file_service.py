"""Generated-file rendering + storage (CLAUDE.md §12, §14) — the ``file_creator``
agent's one code path for turning Markdown-authored content into a real,
downloadable file, and for serving it back.

The LLM always authors the document body in Markdown (or, for ``csv``/``json``,
the raw target text) — ``render()`` converts that into each target format's
native structure. Storage mirrors ``document_service``: bytes go straight to
S3, keyed by a precomputed id so the DB row and the S3 object agree.
"""

from __future__ import annotations

import io
import mimetypes
import re
import unicodedata
import uuid
from html import escape as _html_escape

import structlog
from sqlalchemy.ext.asyncio import AsyncSession

from app.config import settings
from app.core import aws
from app.db.models.generated_file import GeneratedFile
from app.db.repositories.generated_file_repo import GeneratedFileRepository

logger = structlog.get_logger(__name__)

_MIME: dict[str, str] = {
    "md": "text/markdown",
    "txt": "text/plain",
    "csv": "text/csv",
    "json": "application/json",
    "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "pdf": "application/pdf",
    "xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
}


def mime_for(fmt: str, code_ext: str | None = None) -> str:
    if fmt == "code" and code_ext:
        guessed, _ = mimetypes.guess_type(f"file.{code_ext}")
        return guessed or "text/plain"
    return _MIME.get(fmt, "text/plain")


# ─── Markdown → blocks (shared by the docx/pdf/xlsx renderers) ──────────────

_HEADING_RE = re.compile(r"^(#{1,3})\s+(.*)")
_BULLET_RE = re.compile(r"^[-*]\s+(.*)")
_SEPARATOR_RE = re.compile(r"^:?-{1,}:?$")


def _parse_markdown_blocks(md: str) -> list[dict]:
    """A deliberately small Markdown parser — headings, paragraphs, bullet
    lists, and a GFM pipe table. Enough structure for a written-up document;
    not a general Markdown engine."""
    blocks: list[dict] = []
    lines = md.replace("\r\n", "\n").split("\n")
    para_buf: list[str] = []

    def flush_para() -> None:
        text = " ".join(para_buf).strip()
        if text:
            blocks.append({"type": "para", "text": text})
        para_buf.clear()

    i = 0
    while i < len(lines):
        stripped = lines[i].strip()
        if not stripped:
            flush_para()
            i += 1
            continue
        heading = _HEADING_RE.match(stripped)
        if heading:
            flush_para()
            blocks.append(
                {
                    "type": "heading",
                    "level": len(heading.group(1)),
                    "text": heading.group(2).strip(),
                }
            )
            i += 1
            continue
        bullet = _BULLET_RE.match(stripped)
        if bullet:
            flush_para()
            blocks.append({"type": "bullet", "text": bullet.group(1).strip()})
            i += 1
            continue
        if stripped.startswith("|") and stripped.endswith("|"):
            flush_para()
            rows: list[list[str]] = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                cells = [c.strip() for c in lines[i].strip().strip("|").split("|")]
                if not all(_SEPARATOR_RE.match(c) for c in cells):
                    rows.append(cells)
                i += 1
            if rows:
                blocks.append({"type": "table", "rows": rows})
            continue
        para_buf.append(stripped)
        i += 1
    flush_para()
    return blocks


# ─── inline markdown (**bold**, *italic*, `code`, stray <br>) ────────────────
# The LLM writes plain Markdown for the body; both renderers below turn that
# into real formatting instead of printing the markers literally.

_BR_RE = re.compile(r"<br\s*/?>", re.IGNORECASE)
_BOLD_RE = re.compile(r"\*\*([^\n*]+?)\*\*")
_ITALIC_RE = re.compile(r"(?<!\*)\*([^\n*]+?)\*(?!\*)")
_CODE_RE = re.compile(r"`([^`\n]+?)`")

_UNICODE_ASCII_MAP: dict[str, str] = {
    "‘": "'", "’": "'", "‚": "'",  # single quotes
    "“": '"', "”": '"', "„": '"',  # double quotes
    "–": "-", "—": "-", "−": "-", "‐": "-", "‑": "-", "‒": "-", "―": "-",  # dash/hyphen variants
    "…": "...",  # ellipsis
    "•": "-", "·": "-", "●": "-", "‣": "-",  # bullets
    "→": "->", "←": "<-", "↔": "<->",  # arrows
    "≈": "~", "≥": ">=", "≤": "<=", "×": "x", "÷": "/",
    "✓": "v", "✔": "v", "✗": "x", "✘": "x",  # check/cross
    " ": " ", " ": " ", " ": " ", " ": " ",  # odd spaces
}


def _latin1_safe(text: str) -> str:
    """fpdf2's core fonts are Latin-1 only. Swap common "smart" punctuation,
    bullets, arrows and odd spaces for their ASCII equivalent, fold accented
    Latin letters to their plain form (café -> cafe), and only THEN drop
    whatever is still left (CJK, emoji, …) rather than letting it crash or
    silently turn into a wall of "?"."""
    for uni, ascii_ in _UNICODE_ASCII_MAP.items():
        text = text.replace(uni, ascii_)
    text = unicodedata.normalize("NFKD", text)
    text = "".join(c for c in text if not unicodedata.combining(c))
    return text.encode("latin-1", "replace").decode("latin-1")


def _strip_markdown(text: str) -> str:
    """Plain text — markdown emphasis markers removed, not converted (for
    docx headings, whose style is already bold)."""
    text = _BR_RE.sub(" ", text)
    text = _BOLD_RE.sub(r"\1", text)
    text = _ITALIC_RE.sub(r"\1", text)
    return _CODE_RE.sub(r"\1", text)


_INLINE_SPLIT_RE = re.compile(
    r"(\*\*[^\n*]+?\*\*|`[^`\n]+?`|(?<!\*)\*[^\n*]+?\*(?!\*))"
)


def _inline_tokens(text: str) -> list[tuple[str, str]]:
    """``(segment, kind)`` pairs — ``kind`` in ``{"", "b", "i", "code"}`` — for
    building real bold/italic docx runs instead of printing ``**markers**``."""
    tokens: list[tuple[str, str]] = []
    for part in _INLINE_SPLIT_RE.split(_BR_RE.sub(" ", text)):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            tokens.append((part[2:-2], "b"))
        elif part.startswith("`") and part.endswith("`"):
            tokens.append((part[1:-1], "code"))
        elif part.startswith("*") and part.endswith("*"):
            tokens.append((part[1:-1], "i"))
        else:
            tokens.append((part, ""))
    return tokens


def _inline_html(text: str, *, cell: bool = False) -> str:
    """Markdown inline emphasis -> real HTML tags for fpdf2's ``write_html``.
    A stray literal ``<br>`` the model left inside a table cell becomes a
    separator instead of leaking into the rendered text (fpdf2 doesn't support
    a nested ``<br>`` inside ``<td>``); escaping runs BEFORE the emphasis
    conversion so the tags this function introduces survive it."""
    text = _BR_RE.sub("; " if cell else " ", text)
    text = _html_escape(text, quote=False)
    text = _BOLD_RE.sub(r"<b>\1</b>", text)
    text = _ITALIC_RE.sub(r"<i>\1</i>", text)
    return _CODE_RE.sub(r"\1", text)


def _add_docx_runs(paragraph, text: str) -> None:
    for seg, kind in _inline_tokens(text):
        if not seg:
            continue
        run = paragraph.add_run(seg)
        if kind == "b":
            run.bold = True
        elif kind == "i":
            run.italic = True
        elif kind == "code":
            run.font.name = "Consolas"


def _plain_cell(text: str) -> str:
    """Flattened, marker-free text — for renderers with no rich-text cells."""
    return _strip_markdown(_BR_RE.sub("; ", text)).strip()


def _render_docx(content: str) -> bytes:
    from docx import Document as DocxDocument

    doc = DocxDocument()
    for block in _parse_markdown_blocks(content):
        if block["type"] == "heading":
            doc.add_heading(_strip_markdown(block["text"]), level=min(block["level"], 3))
        elif block["type"] == "bullet":
            para = doc.add_paragraph(style="List Bullet")
            _add_docx_runs(para, block["text"])
        elif block["type"] == "table":
            rows = block["rows"]
            table = doc.add_table(rows=0, cols=len(rows[0]))
            table.style = "Light Grid Accent 1"
            for r in rows:
                cells = table.add_row().cells
                for idx, val in enumerate(r):
                    if idx >= len(cells):
                        continue
                    lines = [ln.strip() for ln in _BR_RE.split(val) if ln.strip()] or [""]
                    for j, line in enumerate(lines):
                        para = cells[idx].paragraphs[0] if j == 0 else cells[idx].add_paragraph()
                        _add_docx_runs(para, line)
        else:
            para = doc.add_paragraph()
            _add_docx_runs(para, block["text"])
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _blocks_to_html(blocks: list[dict]) -> str:
    """Parsed blocks -> a small, well-formed HTML subset ``write_html``
    understands: real ``<table border>`` grids (not a flat ``" | "`` string),
    ``<h1>``-``<h3>``, ``<ul><li>``, and inline ``<b>``/``<i>``."""
    parts: list[str] = []
    i = 0
    while i < len(blocks):
        block = blocks[i]
        if block["type"] == "heading":
            level = min(block["level"], 3)
            parts.append(f"<h{level}>{_inline_html(block['text'])}</h{level}>")
            i += 1
        elif block["type"] == "bullet":
            items = []
            while i < len(blocks) and blocks[i]["type"] == "bullet":
                items.append(f"<li>{_inline_html(blocks[i]['text'])}</li>")
                i += 1
            parts.append("<ul>" + "".join(items) + "</ul>")
        elif block["type"] == "table":
            header, *body_rows = block["rows"]
            col_width = f"{100 / max(1, len(header)):.2f}%"
            thead = (
                "<tr>"
                + "".join(
                    f'<th width="{col_width}">{_inline_html(c, cell=True)}</th>' for c in header
                )
                + "</tr>"
            )
            tbody = "".join(
                "<tr>" + "".join(f"<td>{_inline_html(c, cell=True)}</td>" for c in r) + "</tr>"
                for r in body_rows
            )
            parts.append(f'<table border="1"><thead>{thead}</thead><tbody>{tbody}</tbody></table>')
            i += 1
        else:
            parts.append(f"<p>{_inline_html(block['text'])}</p>")
            i += 1
    return "\n".join(parts)


def _render_pdf(content: str) -> bytes:
    from fpdf import FPDF
    from fpdf.fonts import TextStyle

    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=18)
    pdf.set_margins(18, 18, 18)
    pdf.add_page()
    pdf.set_font("Helvetica", size=11)

    body_html = _latin1_safe(_blocks_to_html(_parse_markdown_blocks(content)))
    # h3/h4+ collapse to h3 (_blocks_to_html caps heading level at 3); th/td
    # per-tag styling isn't supported by this fpdf2 version, so the table
    # keeps its default bold-centered header + bordered grid.
    tag_styles = {
        "h1": TextStyle(font_family="helvetica", font_style="B", font_size_pt=19,
                         color=(17, 24, 39), t_margin=0, b_margin=3),
        "h2": TextStyle(font_family="helvetica", font_style="B", font_size_pt=15,
                         color=(31, 41, 55), t_margin=5, b_margin=2),
        "h3": TextStyle(font_family="helvetica", font_style="B", font_size_pt=12,
                         color=(55, 65, 81), t_margin=4, b_margin=2),
        "p": TextStyle(font_family="helvetica", font_size_pt=11, t_margin=0, b_margin=2),
        "li": TextStyle(font_family="helvetica", font_size_pt=11),
    }
    pdf.write_html(
        body_html,
        tag_styles=tag_styles,
        li_prefix_color=(90, 90, 90),
        table_line_separators=True,
    )
    return bytes(pdf.output())


def _render_xlsx(content: str) -> bytes:
    from openpyxl import Workbook
    from openpyxl.styles import Font
    from openpyxl.utils import get_column_letter

    wb = Workbook()
    ws = wb.active
    assert ws is not None  # always true right after Workbook()
    ws.title = "Sheet1"
    table = next((b for b in _parse_markdown_blocks(content) if b["type"] == "table"), None)
    if table:
        for row in table["rows"]:
            ws.append([_plain_cell(c) for c in row])
        for cell in ws[1]:  # bold header row
            cell.font = Font(bold=True)
    else:
        for line in content.strip().splitlines():
            if line.strip():
                ws.append([_plain_cell(line.strip())])
    widths: dict[int, int] = {}
    for row in ws.iter_rows():
        for cell in row:
            if cell.value and cell.column is not None:
                widths[cell.column] = max(widths.get(cell.column, 0), len(str(cell.value)))
    for col, width in widths.items():
        ws.column_dimensions[get_column_letter(col)].width = min(60, max(10, width + 2))
    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


def render(fmt: str, content: str) -> bytes:
    """``content`` (Markdown, or raw CSV/JSON text for those formats) → bytes
    in the target format."""
    if fmt == "docx":
        return _render_docx(content)
    if fmt == "pdf":
        return _render_pdf(content)
    if fmt == "xlsx":
        return _render_xlsx(content)
    # md / txt / csv / json / code — the LLM already wrote the target text directly
    return content.strip().encode("utf-8") + b"\n"


# ─── storage ─────────────────────────────────────────────────────────────────


async def upload(
    db: AsyncSession,
    user_id: uuid.UUID,
    conversation_id: uuid.UUID,
    *,
    filename: str,
    fmt: str,
    data: bytes,
    mime_type: str,
    summary: str | None = None,
) -> GeneratedFile:
    file_id = uuid.uuid4()
    s3_key = f"{user_id}/{conversation_id}/{file_id}/{filename}"
    aws.s3().put_object(
        Bucket=settings.S3_BUCKET_NAME, Key=s3_key, Body=data, ContentType=mime_type
    )
    row = await GeneratedFileRepository(db).add(
        GeneratedFile(
            id=file_id,
            user_id=user_id,
            conversation_id=conversation_id,
            filename=filename,
            format=fmt,
            mime_type=mime_type,
            s3_key=s3_key,
            byte_size=len(data),
            summary=summary,
        )
    )
    logger.info(
        "generated_file_uploaded",
        file_id=str(file_id),
        conversation_id=str(conversation_id),
        format=fmt,
        bytes=len(data),
        s3_key=s3_key,
    )
    return row


async def download_bytes(
    db: AsyncSession, user_id: uuid.UUID, file_id: uuid.UUID
) -> tuple[GeneratedFile, bytes] | None:
    row = await GeneratedFileRepository(db).get_for_user(file_id, user_id)
    if row is None:
        return None
    obj = aws.s3().get_object(Bucket=settings.S3_BUCKET_NAME, Key=row.s3_key)
    data = obj["Body"].read()
    logger.info("generated_file_downloaded", file_id=str(file_id), bytes=len(data))
    return row, data
