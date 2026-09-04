"""file_service — render() per format + upload()/download_bytes() (CLAUDE.md §12, §14)."""

from __future__ import annotations

import uuid

import pytest
from app.services import file_service as fs

_MD = """\
# Report Title

Some intro paragraph about the topic.

## Section

- first point
- second point

| Name | Score |
| --- | --- |
| Alice | 9 |
| Bob | 7 |
"""


def test_render_md_txt_passthrough():
    assert fs.render("md", "# Hi\n\nbody").decode() == "# Hi\n\nbody\n"
    assert fs.render("txt", "just text").decode() == "just text\n"


def test_render_csv_json_passthrough():
    assert fs.render("csv", "a,b\n1,2").decode() == "a,b\n1,2\n"
    assert fs.render("json", '{"a": 1}').decode() == '{"a": 1}\n'


def test_render_docx_is_a_valid_zip():
    data = fs.render("docx", _MD)
    assert data[:2] == b"PK"  # docx is a zip container
    assert len(data) > 100


def test_render_xlsx_is_a_valid_zip():
    data = fs.render("xlsx", _MD)
    assert data[:2] == b"PK"


def test_render_pdf_has_pdf_magic_bytes():
    data = fs.render("pdf", _MD)
    assert data[:4] == b"%PDF"


def test_render_pdf_survives_smart_punctuation():
    # curly quotes / em-dash / ellipsis — must not raise on the Latin-1 core font
    data = fs.render("pdf", "A “quoted” thought — and more…")
    assert data[:4] == b"%PDF"


_TRICKY_MD = """\
# Report Title – 2026

**Bold claim** with *italic* emphasis and a bullet • aside.

| Framework | Notes |
| --- | --- |
| React | 91%<br>86M downloads |
| Vue | Growing |
"""


def _pdf_text(data: bytes) -> str:
    import io as _io

    from pdfminer.high_level import extract_text

    return extract_text(_io.BytesIO(data))


def test_render_pdf_has_no_leaked_markdown_or_html_markers():
    text = _pdf_text(fs.render("pdf", _TRICKY_MD))
    assert "**" not in text
    assert "<br" not in text.lower()
    assert "?" not in text  # no un-transliterated / un-encodable character survived
    assert "React" in text and "91%" in text and "86M downloads" in text


def test_render_pdf_table_is_a_real_bordered_grid_not_flat_text():
    # a flat " | "-joined row would put "React | 91%..." on one text line;
    # a real <table> lays each cell out separately.
    text = _pdf_text(fs.render("pdf", _TRICKY_MD))
    assert "React | 91%" not in text
    assert "Framework" in text and "Notes" in text


def test_render_docx_renders_real_bold_italic_and_splits_br_into_paragraphs():
    import io as _io

    from docx import Document

    doc = Document(_io.BytesIO(fs.render("docx", _TRICKY_MD)))
    body_text = "\n".join(p.text for p in doc.paragraphs)
    assert "**" not in body_text and "<br" not in body_text.lower()

    bold_runs = [r for p in doc.paragraphs for r in p.runs if r.bold and r.text == "Bold claim"]
    italic_runs = [r for p in doc.paragraphs for r in p.runs if r.italic and r.text == "italic"]
    assert bold_runs and italic_runs

    cell = doc.tables[0].rows[1].cells[1]  # React's "Notes" cell
    assert [p.text for p in cell.paragraphs] == ["91%", "86M downloads"]


def test_mime_for_known_and_code():
    assert fs.mime_for("pdf") == "application/pdf"
    assert fs.mime_for("code", "py") == "text/x-python"
    assert fs.mime_for("code", None) == "text/plain"


class _S3:
    calls: list = []
    body: bytes = b""

    def put_object(self, **kw):
        _S3.calls.append(("put", kw["Key"]))

    def get_object(self, **kw):
        _S3.calls.append(("get", kw["Key"]))

        class _Body:
            def read(_self):
                return _S3.body

        return {"Body": _Body()}


class _Repo:
    saved: object | None = None

    def __init__(self, _db): ...

    async def add(self, obj):
        _Repo.saved = obj
        return obj

    async def get_for_user(self, file_id, user_id):
        row = _Repo.saved
        if row is not None and row.id == file_id and row.user_id == user_id:
            return row
        return None


@pytest.fixture(autouse=True)
def _patch(monkeypatch):
    _S3.calls = []
    _S3.body = b"hello"
    _Repo.saved = None
    monkeypatch.setattr(fs.aws, "s3", lambda: _S3())
    monkeypatch.setattr(fs, "GeneratedFileRepository", _Repo)
    monkeypatch.setattr(fs.settings, "S3_BUCKET_NAME", "b")


async def test_upload_puts_to_s3_and_saves_row():
    user_id, conv_id = uuid.uuid4(), uuid.uuid4()
    row = await fs.upload(
        None, user_id, conv_id,
        filename="report.pdf", fmt="pdf", data=b"%PDF-fake",
        mime_type="application/pdf", summary="A short report",
    )
    assert row.filename == "report.pdf"
    assert row.byte_size == len(b"%PDF-fake")
    assert _S3.calls and _S3.calls[0][0] == "put"
    assert str(row.id) in _S3.calls[0][1]


async def test_download_bytes_roundtrip_and_ownership():
    user_id, conv_id = uuid.uuid4(), uuid.uuid4()
    row = await fs.upload(
        None, user_id, conv_id,
        filename="notes.txt", fmt="txt", data=b"hello",
        mime_type="text/plain",
    )
    result = await fs.download_bytes(None, user_id, row.id)
    assert result is not None
    got_row, data = result
    assert got_row.id == row.id
    assert data == b"hello"

    missing = await fs.download_bytes(None, uuid.uuid4(), row.id)
    assert missing is None
